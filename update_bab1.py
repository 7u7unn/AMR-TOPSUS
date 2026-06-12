import docx

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def update_bab1(filename):
    doc = docx.Document(filename)
    
    # Identify headings
    latar_belakang_heading = None
    analisis_masalah_heading = None
    kebutuhan_sistem_heading = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if "Latar Belakang" in text and para.style.name.startswith("Heading"):
            latar_belakang_heading = para
        elif "Analisis Masalah" in text and para.style.name.startswith("Heading"):
            analisis_masalah_heading = para
        elif "Analisis Kebutuhan Sistem" in text and para.style.name.startswith("Heading"):
            kebutuhan_sistem_heading = para

    # Delete everything between Latar Belakang and Analisis Masalah
    delete_mode = False
    for para in doc.paragraphs:
        if para == latar_belakang_heading:
            delete_mode = True
            continue
        if para == analisis_masalah_heading:
            delete_mode = False
            continue
        if delete_mode:
            delete_paragraph(para)
            
    # Delete everything between Analisis Masalah and Analisis Kebutuhan Sistem
    delete_mode = False
    for para in doc.paragraphs:
        if para == analisis_masalah_heading:
            delete_mode = True
            continue
        if para == kebutuhan_sistem_heading:
            delete_mode = False
            continue
        if delete_mode:
            delete_paragraph(para)
            
    # Insert new text after Latar Belakang
    p1 = latar_belakang_heading.insert_paragraph_before("PT. XYZ saat ini membutuhkan sistem pemindahan barang antar gedung di area produksi yang lebih andal dan fleksibel. Sebelumnya, penggunaan sistem pemindahan barang masih memiliki banyak keterbatasan karena kaku dan menyulitkan rekonfigurasi layout pabrik. Oleh karena itu, dikembangkanlah Autonomous Mobile Robot (AMR) berbasis ROS 2 yang mengandalkan navigasi cerdas (Nav2) dan Simultaneous Localization and Mapping (SLAM). Dengan memanfaatkan SLLiDAR A2M7 dan sistem sensor fusion, AMR ini mampu memetakan lingkungan secara mandiri, merencanakan jalur secara dinamis, dan menghindari rintangan (obstacle avoidance) secara real-time, sehingga secara signifikan meningkatkan efisiensi proses logistik internal.")
    latar_belakang_heading._element.addnext(p1._element)

    p2 = analisis_masalah_heading.insert_paragraph_before("1. Keterbatasan Fleksibilitas Navigasi: Sistem pemindahan barang sebelumnya sangat kaku dan rentan terhadap kegagalan operasional jika terjadi perubahan tata letak (layout shifting) pada lantai produksi.\n2. Ketidakmampuan Menghindari Halangan: Tidak adanya sistem obstacle avoidance yang dinamis membuat sistem tidak mampu merespons rintangan tak terduga, yang berujung pada tingginya downtime.\n3. Kebutuhan Manajemen Terpusat: Perlunya sebuah antarmuka pemantauan (web interface) terintegrasi ROS 2 untuk memantau telemetri, mengatur misi navigasi (waypoints), dan mengontrol operasional AMR secara efisien.")
    analisis_masalah_heading._element.addnext(p2._element)
    
    doc.save(filename)
    print("Bab 1 updated.")

if __name__ == '__main__':
    update_bab1(r"d:\MagangUnair26\Topsus - AMR\Preliminary Design Review.docx")
