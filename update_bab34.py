import docx

def delete_paragraph(paragraph):
    p = paragraph._element
    if p.getparent() is not None:
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

def update_bab34(filename):
    doc = docx.Document(filename)
    
    target_keys = [
        "Manajemen Proyek",
        "Work Breakdown Schedule",
        "Alur dan Jadwal Kegiatan",
        "Rencana Pengujian",
        "Rencana Manajemen Risiko",
        "3. LAPORAN KEMAJUAN",
        "4. KESIMPULAN"
    ]
    
    # Store the actual heading elements so we don't rely on equality of wrapper objects
    heading_elements = {}
    
    for para in doc.paragraphs:
        if not para.style.name.startswith("Heading"):
            continue
        text = para.text.strip()
        for key in target_keys:
            if key in text and key not in heading_elements:
                heading_elements[key] = para
                break

    if "Manajemen Proyek" not in heading_elements:
        print("Error: Manajemen Proyek heading not found!")
        return

    # Delete all paragraphs after Manajemen Proyek that are not the target headings
    # We will use the underlying XML element identity
    heading_xmls = [p._element for p in heading_elements.values()]
    manajemen_xml = heading_elements["Manajemen Proyek"]._element
    
    to_delete = []
    past_manajemen = False
    
    for para in doc.paragraphs:
        if para._element == manajemen_xml:
            past_manajemen = True
            
        if not past_manajemen:
            continue
            
        # If it's one of our target headings, keep it
        if para._element in heading_xmls:
            continue
            
        # If it's a heading that starts with our target keys, keep it (just to be safe)
        if para.style.name.startswith("Heading"):
            keep = False
            for key in target_keys:
                if key in para.text:
                    keep = True
                    break
            if keep:
                continue
                
        to_delete.append(para)

    for para in to_delete:
        delete_paragraph(para)
        
    texts = {
        "Manajemen Proyek": "Manajemen proyek mencakup tahap persiapan hardware (Mini PC i7-4600U dan sensor), pengembangan low-level controller (odometri dan TF), implementasi stack navigasi (SLAM dan Nav2), serta pengujian integrasi end-to-end.",
        "Work Breakdown Schedule": "WBS dirancang menjadi beberapa fase utama:\n1. Persiapan Hardware & Low-Level Development: Meliputi integrasi Mini PC i7-4600U, pengembangan jembatan odometri dari motor driver ZLAC1525, publikasi TF base_link, dan fusi sensor IMU N100 dengan odometri roda.\n2. Implementasi ROS 2 & High-Level Navigation: Meliputi instalasi lingkungan Ubuntu 22.04 dan ROS 2 Humble, tuning SLAM Toolbox untuk pembuatan peta Occupancy Grid, dan konfigurasi Navigation 2 (Smac Planner & MPPI Controller).\n3. Integrasi Sistem & Web Interface: Meliputi penyelesaian REST API dan WebSocket untuk kontrol telemetri jarak jauh.\n4. Pengujian Performa: Validasi kinerja end-to-end dari seluruh sistem.",
        "Alur dan Jadwal Kegiatan": "Alur kegiatan dimulai dari pengujian low-level driver pada Mini PC i7-4600U, kalibrasi sensor, serta pengembangan node odometri dan TF. Setelah pondasi robot solid, dilanjutkan dengan pengaturan parameter Nav2. Setiap tahap memiliki milestone pengujian fungsional sebelum digabungkan ke dalam peluncuran sistem penuh (bringup).",
        "Rencana Pengujian": "Rencana pengujian meliputi kalibrasi odometri aktual yang divalidasi dengan pergerakan translasi dan rotasi murni, uji fusi sensor melalui Robot Localization (EKF), validasi akurasi peta menggunakan LiDAR A2M7, serta evaluasi kecepatan dan presisi penghindaran rintangan dinamis oleh Nav2. Pengujian integrasi Web Interface juga dilakukan secara komprehensif.",
        "Rencana Manajemen Risiko": "Risiko utama dan strategi mitigasinya meliputi:\n1. Gangguan Komunikasi Node ROS 2: Dapat terjadi akibat beban komputasi CPU. Mitigasi dilakukan dengan mengoptimalkan rate publikasi dari node pada Mini PC i7-4600U.\n2. Masalah Sinkronisasi TF Tree: Transformasi yang tertunda dapat merusak navigasi. Mitigasi berupa konfigurasi waktu menggunakan NTP/chrony untuk menjaga stabilitas stempel waktu ROS 2.\n3. Akumulasi Drift Odometri: Mitigasi melalui tuning cermat pada node EKF menggunakan presisi dari IMU N100.",
        "3. LAPORAN KEMAJUAN": "AMR Arya telah berhasil diintegrasikan dengan arsitektur ROS 2 Humble yang terpusat pada Mini PC i7-4600U. Modul low-level (seperti publikasi odometri motor dan TF) kini beroperasi stabil, memungkinkan integrasi high-level (SLAM Toolbox dan Nav2) untuk menjalankan navigasi otonom di lingkungan simulasi maupun lingkungan uji nyata. Modul arya_web_interface juga telah sukses dikendalikan sebagai pusat kendali jarak jauh.",
        "4. KESIMPULAN": "Peralihan desain AMR Arya menuju framework ROS 2 dan adopsi Mini PC i7-4600U memberikan lompatan keandalan operasional yang substansial. Kemampuan untuk membangun layer low-level yang kokoh sebagai pondasi stack navigasi tinggi membuktikan bahwa AMR Arya siap menggantikan peran AGV konvensional di PT. XYZ, dengan menawarkan kebebasan manuver tanpa batas fisik dan penghindaran rintangan otomatis (obstacle avoidance)."
    }

    for key, text in texts.items():
        if key in heading_elements:
            heading = heading_elements[key]
            paras = text.split('\n')
            curr = heading
            for p_text in paras:
                p_text = p_text.strip()
                if p_text:
                    new_p = heading.insert_paragraph_before(p_text)
                    curr._element.addnext(new_p._element)
                    curr = new_p

    doc.save(filename)
    print("Bab 3 and 4 successfully updated.")

if __name__ == '__main__':
    update_bab34(r"d:\MagangUnair26\Topsus - AMR\Preliminary Design Review.docx")
