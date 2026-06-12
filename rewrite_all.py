import docx

def delete_paragraph(paragraph):
    p = paragraph._element
    if p.getparent() is not None:
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

def rewrite_all(filename):
    doc = docx.Document(filename)
    
    target_headings = [
        "Ringkasan Isi Dokumen",
        "Tujuan Penulisan Dokumen",
        "Referensi",
        "Daftar Singkatan",
        "Latar Belakang",
        "Informasi Pendukung masalah",
        "Analisis Masalah",
        "Analisis Kebutuhan Sistem",
        "Kebutuhan Desain Sistem",
        "Kebutuhan Performansi Sistem",
        "Arsitektur Desain Sistem",
        "Alternatif Solusi",
        "Analisis Solusi",
        "Solusi yang Dipilih",
        "Karakteristik Produk",
        "Manajemen Proyek",
        "Work Breakdown Schedule",
        "Alur dan Jadwal Kegiatan",
        "Rencana Pengujian",
        "Rencana Manajemen Risiko",
        "LAPORAN KEMAJUAN",
        "KESIMPULAN"
    ]
    
    # Locate all target headings based on XML element
    heading_elements = {}
    for para in doc.paragraphs:
        if not para.style.name.startswith("Heading"):
            continue
        text = para.text.strip()
        for key in target_headings:
            if key.lower() in text.lower() and key not in heading_elements:
                heading_elements[key] = para
                break

    # To cleanly delete contents UNDER these headings, we iterate through all paragraphs.
    # If a paragraph is a heading in our target list, we mark it as the current active heading.
    # If a paragraph is a heading NOT in our target list (e.g. "1. PENGANTAR"), we unset active heading.
    # If a paragraph is normal text and we are under an active heading, we delete it.
    
    heading_xmls = [p._element for p in heading_elements.values()]
    
    to_delete = []
    active_heading_key = None
    
    for para in doc.paragraphs:
        if para._element in heading_xmls:
            # Find which key it is
            for k, v in heading_elements.items():
                if v._element == para._element:
                    active_heading_key = k
                    break
            continue
            
        if para.style.name.startswith("Heading"):
            # It's a heading but not one of our targets, so we stop deleting
            active_heading_key = None
            continue
            
        if active_heading_key is not None:
            # Don't delete table of contents or empty spaces that might be important?
            # Table of contents is usually a special field, but its lines might show up as paragraphs.
            # To be safe, we'll only delete if it has text or if we just want to wipe it.
            to_delete.append(para)
            
    for para in to_delete:
        delete_paragraph(para)
        
    texts = {
        "Ringkasan Isi Dokumen": "Dokumen ini merangkum rancangan sistem Autonomous Mobile Robot (AMR) Arya berbasis ROS 2 Humble untuk aplikasi navigasi indoor di fasilitas industri PT. XYZ. Sistem ini dikendalikan oleh Mini PC i7-4600U yang terintegrasi dengan sensor SLLiDAR A2M7, CCF-LAS6-4M laser sensor, IMU N100, dan motor driver ZLAC1525. AMR ini difokuskan pada kemampuan pemetaan mandiri (SLAM), lokalisasi presisi (AMCL), dan navigasi penghindaran rintangan dinamis (Nav2). Arsitektur ini juga menyertakan Web Interface untuk telemetri dan kontrol jarak jauh. Dokumen ini juga membahas manajemen proyek, jadwal kegiatan, dan rencana pengujian yang komprehensif.",
        "Tujuan Penulisan Dokumen": "1. Mendokumentasikan rancangan arsitektur sistem AMR Arya berbasis ROS 2.\n2. Mengidentifikasi kebutuhan fungsional dan performansi sistem navigasi tingkat industri.\n3. Menyusun rencana manajemen proyek, pengujian operasional, dan mitigasi risiko.",
        "Referensi": "Referensi mencakup dokumentasi resmi ROS 2 Humble, panduan integrasi Nav2, dokumentasi SLAM Toolbox, serta datasheet komponen perangkat keras seperti SLLiDAR A2M7 dan Motor Driver ZLAC1525.",
        "Daftar Singkatan": "AMR : Autonomous Mobile Robot\nROS : Robot Operating System\nSLAM : Simultaneous Localization and Mapping\nNav2 : Navigation 2\nIMU : Inertial Measurement Unit\nEKF : Extended Kalman Filter\nAPI : Application Programming Interface",
        "Latar Belakang": "PT. XYZ saat ini membutuhkan sistem logistik internal antar gedung yang andal. Sebelumnya, sistem Automated Guided Vehicle (AGV) konvensional dengan pita magnetik yang kaku menyebabkan downtime operasional yang tinggi saat terjadi perubahan tata letak fasilitas. Oleh karena itu, AMR Arya berbasis ROS 2 dikembangkan untuk memberikan fleksibilitas tinggi dengan mengandalkan pemetaan berbasis LiDAR dan sistem navigasi cerdas yang memungkinkan rekonfigurasi rute secara dinamis tanpa mengubah infrastruktur fisik.",
        "Informasi Pendukung masalah": "Sistem AGV konvensional rentan terhadap keausan jalur magnetik di lantai pabrik dan memiliki ketidakmampuan merespons rintangan tak terduga. Hal ini memaksa terjadinya intervensi operator yang menurunkan efisiensi waktu distribusi barang.",
        "Analisis Masalah": "1. Infrastruktur fisik kaku: Bergantung sepenuhnya pada jalur magnetik.\n2. Tidak ada obstacle avoidance: Rentan terhadap rintangan dinamis di lantai pabrik, menyebabkan downtime.\n3. Kurangnya manajemen terpusat: Tidak adanya antarmuka kontrol jarak jauh untuk mengoordinasikan armada.",
        "Analisis Kebutuhan Sistem": "Kebutuhan sistem diturunkan dari tujuan AMR untuk beroperasi secara mandiri, memetakan pabrik, dan merencanakan rute secara real-time. Sistem harus tangguh secara komputasi sehingga arsitektur ROS 2 yang terpusat pada Mini PC i7-4600U dipilih sebagai fondasi operasi.",
        "Kebutuhan Desain Sistem": "Hardware mencakup Mini PC i7-4600U, SLLiDAR A2M7, IMU N100, CCF-LAS6-4M laser sensor, dan Motor Driver ZLAC1525. Software mencakup Ubuntu 22.04, ROS 2 Humble, SLAM Toolbox, Nav2, Robot Localization (EKF), dan Web Interface.",
        "Kebutuhan Performansi Sistem": "Sistem dituntut mampu melakukan pembaruan odometri resolusi tinggi dengan EKF, navigasi dengan replanning cepat menggunakan Smac Planner dan MPPI Controller, dan telemetri jarak jauh yang ber-latency rendah.",
        "Arsitektur Desain Sistem": "Arsitektur dibagi menjadi empat lapisan:\n1. Lapisan Web Interface: Akses telemetri dan kontrol jarak jauh.\n2. Lapisan Navigation Stack: Menjalankan eksekusi Nav2 dan Behavior Tree.\n3. Lapisan Sensor Fusion & Drivers: Memproses fusi sensor IMU dan LiDAR melalui EKF.\n4. Lapisan Motor Control: Mengonversi perintah kecepatan menjadi aksi pada Motor Driver ZLAC1525.",
        "Alternatif Solusi": "Pendekatan dalam industri umumnya terbagi menjadi navigasi berbasis infrastruktur statis (seperti garis atau rel) dan navigasi otonom dinamis. Navigasi statis memiliki biaya pemeliharaan tinggi dan fleksibilitas rendah.",
        "Analisis Solusi": "Pendekatan otonom dinamis menggunakan ROS 2 dan Nav2 terbukti memberikan solusi penghindaran halangan secara real-time. Penggunaan Mini PC i7-4600U dipilih karena kemampuannya secara lancar memproses algoritma probabilistik berat seperti AMCL dan EKF tanpa bottleneck komputasi.",
        "Solusi yang Dipilih": "Solusi akhir menggunakan sistem navigasi otonom berbasis graf (Nav2) dengan fusi sensor ekstensif (LiDAR & IMU) yang dijalankan sepenuhnya pada Mini PC i7-4600U.",
        "Karakteristik Produk": "Karakteristik produk adalah sistem kendali differential drive murni dengan operasi nirkabel penuh, obstacle avoidance dinamis, serta kemampuan dipantau melalui antarmuka web tersentralisasi.",
        "Manajemen Proyek": "Manajemen proyek mencakup tahapan penyiapan hardware komputasi, pengembangan driver low-level, implementasi algoritma navigasi cerdas, serta pengujian sistem.",
        "Work Breakdown Schedule": "WBS dirancang menjadi:\n1. Persiapan Hardware & Driver: Integrasi Mini PC i7-4600U, odometri ZLAC1525, dan TF base_link.\n2. Implementasi ROS 2 & High-Level Navigation: Instalasi Ubuntu 22.04, SLAM Toolbox, dan Nav2.\n3. Integrasi Web Interface.\n4. Pengujian Performa Keseluruhan.",
        "Alur dan Jadwal Kegiatan": "Kegiatan dimulai dari pengujian low-level driver (odometry dan IMU), kalibrasi sensor, diakhiri dengan pengaturan parameter Nav2 dan bringup secara keseluruhan.",
        "Rencana Pengujian": "Meliputi kalibrasi odometri, pengujian presisi EKF, pengujian akurasi peta SLAM, validasi obstacle avoidance dinamis, serta uji ketahanan koneksi antarmuka web.",
        "Rencana Manajemen Risiko": "1. Beban CPU tinggi: Mitigasi dengan optimasi node ROS 2 di i7-4600U.\n2. Isu sinkronisasi waktu: Mitigasi dengan setup chrony/NTP.\n3. Drift odometri: Mitigasi dengan tuning EKF dan IMU.",
        "LAPORAN KEMAJUAN": "AMR Arya telah berhasil diintegrasikan dengan arsitektur ROS 2 Humble. Modul low-level beroperasi stabil, sementara SLAM Toolbox dan Nav2 telah berfungsi mengarahkan robot pada lingkungan uji. Antarmuka web (arya_web_interface) juga telah fungsional sebagai pusat komando.",
        "KESIMPULAN": "Peralihan menuju AMR Arya dengan kapabilitas komputasi Mini PC i7-4600U dan arsitektur ROS 2 memberikan keandalan navigasi yang dinamis. Desain ini menghapus keterbatasan infrastruktur fisik yang selama ini membebani operasional AGV lama di PT. XYZ."
    }

    for key, text in texts.items():
        if key in heading_elements:
            heading = heading_elements[key]
            paras = text.split('\n')
            curr = heading
            for p_text in paras:
                p_text = p_text.strip()
                if p_text:
                    new_p = heading.insert_paragraph_before(p_text)
                    curr._element.addnext(new_p._element)
                    curr = new_p

    doc.save(filename)
    print("All sections completely rewritten and sanitized.")

if __name__ == '__main__':
    rewrite_all(r"d:\MagangUnair26\Topsus - AMR\Preliminary Design Review.docx")
