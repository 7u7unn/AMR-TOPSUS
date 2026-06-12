import docx

def delete_paragraph(paragraph):
    p = paragraph._element
    if p.getparent() is not None:
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

def update_bab2(filename):
    doc = docx.Document(filename)
    
    headings = {
        "Analisis Kebutuhan Sistem": None,
        "Kebutuhan Desain Sistem": None,
        "Kebutuhan Performansi Sistem": None,
        "Arsitektur Desain Sistem": None,
        "Alternatif Solusi": None,
        "Analisis Solusi": None,
        "Solusi yang Dipilih": None,
        "Karakteristik Produk": None,
        "Manajemen Proyek": None,
    }
    
    # Identify headings
    for para in doc.paragraphs:
        text = para.text.strip()
        if not para.style.name.startswith("Heading"): continue
        
        for key in headings.keys():
            if key in text:
                headings[key] = para
                break

    # We want to clear contents under each heading from Analisis Kebutuhan Sistem up to Manajemen Proyek
    # and then insert new texts.
    
    # Ordered keys to iterate
    keys_in_order = [
        "Analisis Kebutuhan Sistem",
        "Kebutuhan Desain Sistem",
        "Kebutuhan Performansi Sistem",
        "Arsitektur Desain Sistem",
        "Alternatif Solusi",
        "Analisis Solusi",
        "Solusi yang Dipilih",
        "Karakteristik Produk",
        "Manajemen Proyek"
    ]
    
    for i in range(len(keys_in_order) - 1):
        start_heading = headings[keys_in_order[i]]
        end_heading = headings[keys_in_order[i+1]]
        
        if not start_heading: continue
        
        # Collect paragraphs between start_heading and end_heading
        to_delete = []
        delete_mode = False
        for para in doc.paragraphs:
            if para == start_heading:
                delete_mode = True
                continue
            if para == end_heading:
                delete_mode = False
                break
            if delete_mode:
                # keep images/tables if they are not in paragraphs? No, python-docx doc.paragraphs only yields paragraphs.
                # Just delete paragraphs
                to_delete.append(para)
                
        for para in to_delete:
            delete_paragraph(para)
            
    # Now insert texts
    texts = {
        "Analisis Kebutuhan Sistem": "Analisis kebutuhan sistem mencakup kebutuhan fungsional, non-fungsional, dan performansi. Kebutuhan ini diturunkan dari tujuan AMR untuk memetakan lingkungan pabrik secara presisi, merencanakan jalur yang aman secara mandiri, menghindari rintangan statis dan dinamis, serta menjaga stabilitas kontrol pergerakan. Keseluruhan sistem dibangun di atas arsitektur ROS 2 untuk memastikan keandalan, modularitas, dan interoperabilitas perangkat keras tingkat industri.",
        "Kebutuhan Desain Sistem": "Kebutuhan desain sistem mencakup pemilihan komponen perangkat keras standar industri dan kerangka perangkat lunak yang andal. Hardware meliputi: SLLiDAR A2M7 (untuk pemetaan ruang 2D), CCF-LAS6-4M laser sensor (untuk halangan), Modul IMU N100 (orientasi), Motor Driver ZLAC1525, dan Modul I/O WaveShare Modbus.\n\nDari sisi perangkat lunak, sistem menggunakan arsitektur ROS 2 Humble pada Ubuntu 22.04. Perangkat lunak mencakup driver sensor, SLAM Toolbox untuk peta Occupancy Grid, Robot Localization menggunakan filter EKF, AMCL untuk lokalisasi, serta Nav2 (Smac Planner dan MPPI Controller). Sistem juga mengintegrasikan antarmuka web melalui REST API dan WebSocket untuk kendali jarak jauh.",
        "Kebutuhan Performansi Sistem": "Kebutuhan performansi difokuskan pada stabilitas estimasi pose dan kecepatan respons navigasi. AMR dituntut mampu memproses pembaruan odometri dengan resolusi tinggi yang distabilkan melalui fusi sensor (LiDAR dan IMU). Framework Nav2 harus sanggup melakukan deteksi rintangan dinamis secara real-time dan menghitung ulang rute (replanning) tanpa jeda yang signifikan agar distribusi logistik internal tidak terhambat.",
        "Arsitektur Desain Sistem": "Arsitektur desain sistem disusun secara terpusat memanfaatkan ekosistem ROS 2. Arsitektur ini terbagi menjadi empat lapisan utama:\n1. Lapisan Web Interface: Menyediakan REST API dan WebSocket untuk kendali misi.\n2. Lapisan Navigation Stack (Nav2): Mengatur Planner Server, Controller Server, dan Behavior Server.\n3. Lapisan Sensor Fusion & Drivers: Memproses data SLLiDAR, sensor CCF, dan IMU N100 melalui EKF.\n4. Lapisan Motor Control: Mengeksekusi perintah Nav2 menjadi putaran roda aktual melalui Motor Driver ZLAC1525.",
        "Alternatif Solusi": "Penggunaan mikrokontroler atau komputer berdaya rendah (seperti prototipe awal) terbukti rawan terhadap isu kehabisan memori (Out of Memory) ketika dihadapkan pada komputasi berat seperti SLAM dan Nav2.",
        "Analisis Solusi": "Alternatif yang diterapkan saat ini adalah memanfaatkan PC berbasis Ubuntu dengan RAM yang memadai untuk menjalankan framework ROS 2 secara mulus. Ini memberikan fondasi kuat untuk pemrosesan navigasi canggih tanpa mengorbankan performa.",
        "Solusi yang Dipilih": "Solusi berpusat pada integrasi ROS 2 Humble. AMR mampu menjalankan pemetaan dinamis (SLAM), lokalisasi presisi (AMCL), dan navigasi otonom terintegrasi (Nav2) secara end-to-end tanpa memerlukan panduan jalur fisik.",
        "Karakteristik Produk": "Karakteristik utama AMR yang dikembangkan adalah sistem kendali differential drive murni yang sanggup beroperasi merespons rintangan secara cerdas, memetakan secara otonom, dan dipantau serta dikontrol kinerjanya dari antarmuka web tersentralisasi."
    }

    for key, text in texts.items():
        heading = headings.get(key)
        if heading:
            # handle multiple paragraphs
            paras = text.split('\n')
            curr = heading
            for p_text in paras:
                p_text = p_text.strip()
                if p_text:
                    new_p = heading.insert_paragraph_before(p_text)
                    curr._element.addnext(new_p._element)
                    curr = new_p

    doc.save(filename)
    print("Bab 2 updated.")

if __name__ == '__main__':
    update_bab2(r"d:\MagangUnair26\Topsus - AMR\Preliminary Design Review.docx")
